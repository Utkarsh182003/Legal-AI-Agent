import os
import asyncio
import uuid
from typing import Optional, Dict, List, Any
from utils.llm_utils import load_api_key_from_env

load_api_key_from_env()

from pydantic_ai.models.google import GoogleModel
from agents.document_reader import DocumentReaderAgent
from agents.information_extractor import InformationExtractionAgent
from agents.compliance_analyzer import ComplianceAnalyzerAgent
from agents.knowledge_graph_agent import KnowledgeGraphAgent
from agents.rag_agent import RAGAgent # RAGAgent class still imported
from models.document_models import DocumentInput, DocumentAnalysisResult, DocumentContent, Paragraph, Sentence, ComplianceFinding, Node, Edge, KnowledgeGraph, RAGResponse
from utils.redis_cache import RedisCache

DOCUMENTS_DIR = "documents"
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Singleton instances for Redis cache and LLM model
_redis_cache_instance: Optional[RedisCache] = None
_llm_model_instance: Optional[GoogleModel] = None

def get_redis_cache() -> RedisCache:
    """Returns a singleton RedisCache instance."""
    global _redis_cache_instance
    if _redis_cache_instance is None:
        _redis_cache_instance = RedisCache()
    return _redis_cache_instance

def get_llm_model() -> GoogleModel:
    """Returns a singleton LLM model instance."""
    global _llm_model_instance
    if _llm_model_instance is None:
        _llm_model_instance = GoogleModel("gemini-1.5-flash")
    return _llm_model_instance


def create_dummy_docx(file_path: str):
    """Creates a dummy DOCX file for testing."""
    from docx import Document
    doc = Document()
    doc.add_heading('Sample Lease Agreement', level=1)
    doc.add_paragraph('This Lease Agreement ("Agreement") is made and entered into on this 1st day of January, 2025 ("Effective Date"), by and between Lessor, ABC Corp, located at 123 Main St, Anytown, and Lessee, XYZ LLC, located at 456 Oak Ave, Somewhere. The rent shall be $1,500 USD per month, payable on the 5th day of each month. This Agreement shall be governed by and construed in accordance with the laws of the State of New York. Page 1 of 2.')
    doc.add_paragraph('Any dispute arising out of or in connection with this Agreement shall be subject to the exclusive jurisdiction of the courts of New York. This clause acts as an indemnification for certain events. Force Majeure: Neither party shall be liable for any failure to perform its obligations where such failure is as a result of Acts of God, war, or other circumstances beyond the party\'s reasonable control. Notice period is 30 days.')
    doc.add_page_break()
    doc.add_paragraph('This is the second page of the document. This is boilerplate text. Confidentiality: All information exchanged hereunder is confidential for 5 years. This term shall be binding for a period of five (5) years from the Effective Date. This confidentiality clause is very strict.')
    doc.save(file_path)

def create_dummy_pdf(file_path: str):
    """Creates a dummy PDF file for testing using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()

    text = """
    POLICY DOCUMENT: DATA PRIVACY

    Effective Date: March 15, 2024.

    This Data Privacy Policy ("Policy") outlines the commitment of our organization to protecting the privacy and security of personal data.
    1. Data Collection: We collect personal data solely for legitimate business purposes. This section emphasizes data minimization.
    2. Data Usage: Personal data will only be used for the purposes for which it was collected.
    3. Data Disclosure: We do not disclose personal data to third parties without explicit consent, except as required by law.
    4. Security Measures: We implement robust technical and organizational measures to protect personal data from unauthorized access, alteration, disclosure, or destruction.

    Compliance with GDPR and CCPA is paramount. This policy may be terminated by either party with 90 days notice.
    Page 1 of 1.
    """
    page.insert_text((50, 50), text, fontsize=12)
    doc.save(file_path)
    doc.close()

def create_dummy_txt(file_path: str):
    """Creates a dummy TXT file for testing."""
    text = """
    Simple Legal Note

    Date: 2023-11-20

    Regarding the recent incident, a fine of $250.00 USD was imposed for a breach of conduct.
    Further actions may be taken.
    This document serves as notification. No explicit termination clause or confidentiality clause here.
    """
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)

async def analyze_document_pipeline_core(
    file_path: str, 
    llm_model: GoogleModel, 
    redis_cache: RedisCache,
    rag_agent_instance: RAGAgent 
) -> Optional[DocumentAnalysisResult]:
    """
    Orchestrates the full document analysis pipeline.
    Agents are instantiated *within* this function.
    """
    print(f"\n--- Processing: {os.path.basename(file_path)} ---")
    try:
        reader_agent = DocumentReaderAgent() 
        extractor_agent = InformationExtractionAgent(model=llm_model, cache=redis_cache)
        compliance_agent = ComplianceAnalyzerAgent(model=llm_model, cache=redis_cache)
        knowledge_graph_agent = KnowledgeGraphAgent(model=llm_model, cache=redis_cache)

        document_content: DocumentContent = reader_agent.run(DocumentInput(file_path=file_path))
        print(f"Successfully read and preprocessed '{document_content.file_name}'.")

        analysis_result: DocumentAnalysisResult = await extractor_agent.run(document_content)
        print(f"Successfully extracted information from '{analysis_result.file_name}'.")

        analysis_result = await compliance_agent.run(analysis_result)
        print(f"Compliance analysis complete for '{analysis_result.file_name}'.")

        analysis_result = await knowledge_graph_agent.run(analysis_result)
        print(f"Knowledge Graph construction complete for '{analysis_result.file_name}'.")

        rag_agent_instance.load_document_context(analysis_result)

        print(f"Full analysis pipeline complete for '{analysis_result.file_name}'.")
        return analysis_result

    except Exception as e:
        print(f"An error occurred during document analysis for {file_path}: {e}")
        return None

async def run_rag_query_core(query: str, rag_agent_instance: RAGAgent) -> RAGResponse:
    """
    Runs a RAG query using the provided RAG agent instance.
    """
    return await rag_agent_instance.run(query)

async def setup_dummy_documents():
    """
    Sets up dummy documents.
    """
    dummy_docx_path = os.path.join(DOCUMENTS_DIR, "sample_agreement.docx")
    dummy_pdf_path = os.path.join(DOCUMENTS_DIR, "sample_policy.pdf")
    dummy_txt_path = os.path.join(DOCUMENTS_DIR, "sample_note.txt")

    if not os.path.exists(dummy_docx_path):
        create_dummy_docx(dummy_docx_path)
        print(f"Dummy DOCX file created at: {dummy_docx_path}")
    if not os.path.exists(dummy_pdf_path):
        create_dummy_pdf(dummy_pdf_path)
        print(f"Dummy PDF file created at: {dummy_pdf_path}")
    if not os.path.exists(dummy_txt_path):
        create_dummy_txt(dummy_txt_path)
        print(f"Dummy TXT file created at: {dummy_txt_path}")

if __name__ == "__main__":
    print("Running backend service setup and a single document analysis for testing...")
    
    llm = get_llm_model()
    cache = get_redis_cache()
    
    temp_rag_agent = RAGAgent(model=llm, cache=cache)

    asyncio.run(setup_dummy_documents())
    
    test_doc_path = os.path.join(DOCUMENTS_DIR, "sample_agreement.docx")
    analyzed_result = asyncio.run(analyze_document_pipeline_core(test_doc_path, llm, cache, temp_rag_agent))
    
    if analyzed_result:
        print("\n--- Test Document Analysis Summary (from backend_service.py direct run) ---")
        print(f"Document Type: {analyzed_result.metadata.document_type}")
        print(f"Title: {analyzed_result.metadata.title}")
        print(f"Effective Date: {analyzed_result.metadata.effective_date}")
        print(f"Jurisdiction: {analyzed_result.metadata.jurisdiction}")
        print(f"Analysis Summary: {analyzed_result.analysis_summary}")
        print(f"Number of Parties: {len(analyzed_result.extracted_parties)}")
        print(f"Number of Compliance Findings: {len(analyzed_result.compliance_findings)}")
        if analyzed_result.knowledge_graph:
            print(f"KG Nodes: {len(analyzed_result.knowledge_graph.nodes)}, Edges: {len(analyzed_result.knowledge_graph.edges)}")

        print("\n--- Test RAG Query (from backend_service.py direct run) ---")
        test_query = "Who is the Lessor?"
        rag_response = asyncio.run(run_rag_query_core(test_query, temp_rag_agent))
        print(f"Query: {test_query}")
        print(f"Answer: {rag_response.answer} (Confidence: {rag_response.confidence})")
        if rag_response.relevant_snippets:
            print(f"Relevant Snippets: {rag_response.relevant_snippets[0][:100]}...")
        if rag_response.source_nodes:
            print(f"Source Nodes: {rag_response.source_nodes}")

    print("\nBackend service test complete. Use 'streamlit run app.py' to launch the UI.")

