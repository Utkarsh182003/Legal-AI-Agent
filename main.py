# main.py (UPDATED - Improved dummy_txt for Monetary Value Reason)

import os
import asyncio
import uuid
from utils.llm_utils import load_api_key_from_env
from typing import Optional

# Load environment variables from .env file at the very beginning
load_api_key_from_env()

# Now import pydantic_ai's GoogleModel
from pydantic_ai.models.google import GoogleModel

from agents.document_reader import DocumentReaderAgent
from agents.information_extractor import InformationExtractionAgent
from agents.compliance_analyzer import ComplianceAnalyzerAgent
from agents.knowledge_graph_agent import KnowledgeGraphAgent
from agents.rag_agent import RAGAgent
from models.document_models import DocumentInput, DocumentAnalysisResult, DocumentContent, Paragraph, Sentence, ComplianceFinding, Node, Edge, KnowledgeGraph, RAGResponse
from utils.redis_cache import RedisCache

# Ensure the 'documents' directory exists
DOCUMENTS_DIR = "documents"
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Initialize Redis Cache (optional, only if Redis server is running)
redis_cache = RedisCache()

# Initialize LLM Model for comprehensive extraction and RAG with Google Gemini
llm_model = GoogleModel("gemini-1.5-flash")


def create_dummy_docx(file_path: str):
    """Creates a dummy DOCX file for testing."""
    from docx import Document
    doc = Document()
    doc.add_heading('Sample Lease Agreement', level=1)
    doc.add_paragraph('This Lease Agreement ("Agreement") is made and entered into on this 1st day of January, 2025 ("Effective Date"), by and between Lessor, ABC Corp, located at 123 Main St, Anytown, and Lessee, XYZ LLC, located at 456 Oak Ave, Somewhere. The rent shall be $1,500 USD per month, payable on the 5th day of each month. This Agreement shall be governed by and construed in accordance with the laws of the State of New York. Page 1 of 2.')
    doc.add_paragraph('Any dispute arising out of or in connection with this Agreement shall be subject to the exclusive jurisdiction of the courts of New York. This clause acts as an indemnification for certain events. Force Majeure: Neither party shall be liable for any failure to perform its obligations where such failure is as a result of Acts of God, war, or other circumstances beyond the party\'s reasonable control. Notice period is 30 days.')
    doc.add_page_break()
    doc.add_paragraph('This is the second page of the document. This is boilerplate text. Confidentiality: All information exchanged hereunder is confidential for 5 years. This term shall be binding for a period of five (5) years from the Effective Date. This confidentiality clause is very strict.')
    doc.save(file_path)
    print(f"Dummy DOCX file created at: {file_path}")

def create_dummy_pdf(file_path: str):
    """Creates a dummy PDF file for testing using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()

    text = """
    POLICY DOCUMENT: DATA PRIVACY

    Effective Date: March 15, 2024.

    This Data Privacy Policy ("Policy") outlines the commitment of our organization to protecting the privacy and security of personal data.
    1. Data Collection: We collect personal data solely for legitimate business purposes. This section emphasizes data minimization.
    2. Data Usage: Personal data will only be used for the purposes for which it was collected.
    3. Data Disclosure: We do not disclose personal data to third parties without explicit consent, except as required by law.
    4. Security Measures: We implement robust technical and organizational measures to protect personal data from unauthorized access, alteration, disclosure, or destruction.

    Compliance with GDPR and CCPA is paramount. This policy may be terminated by either party with 90 days notice.
    Page 1 of 1.
    """
    page.insert_text((50, 50), text, fontsize=12)
    doc.save(file_path)
    doc.close()
    print(f"Dummy PDF file created at: {file_path}")

def create_dummy_txt(file_path: str):
    """Creates a dummy TXT file for testing."""
    text = """
    Simple Legal Note

    Date: 2023-11-20

    Regarding the recent incident, a fine of $250.00 USD was imposed for a breach of conduct.
    Further actions may be taken.
    This document serves as notification. No explicit termination clause or confidentiality clause here.
    """
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Dummy TXT file created at: {file_path}")

async def analyze_document(file_path: str, rag_agent: RAGAgent) -> Optional[DocumentAnalysisResult]:
    """
    Orchestrates the document reading and information extraction process.
    """
    reader_agent = DocumentReaderAgent()
    extractor_agent = InformationExtractionAgent(model=llm_model, cache=redis_cache)
    compliance_agent = ComplianceAnalyzerAgent(model=llm_model, cache=redis_cache)
    knowledge_graph_agent = KnowledgeGraphAgent(model=llm_model, cache=redis_cache)

    print(f"\n--- Processing: {os.path.basename(file_path)} ---")
    try:
        document_content: DocumentContent = reader_agent.run(DocumentInput(file_path=file_path))
        print(f"Successfully read and preprocessed '{document_content.file_name}'.")

        analysis_result: DocumentAnalysisResult = await extractor_agent.run(document_content)
        print(f"Successfully extracted information from '{analysis_result.file_name}'.")

        analysis_result = await compliance_agent.run(analysis_result)
        print(f"Compliance analysis complete for '{analysis_result.file_name}'.")

        analysis_result = await knowledge_graph_agent.run(analysis_result)
        print(f"Knowledge Graph construction complete for '{analysis_result.file_name}'.")

        rag_agent.load_document_context(analysis_result)


        print("\n--- Document Analysis Summary ---")
        print(f"Document ID: {analysis_result.document_id}")
        print(f"File Name: {analysis_result.file_name}")
        print(f"Document Type: {analysis_result.metadata.document_type}")
        print(f"Title: {analysis_result.metadata.title}")
        print(f"Effective Date: {analysis_result.metadata.effective_date}")
        print(f"Jurisdiction: {analysis_result.metadata.jurisdiction}")
        print("\nParties:")
        for party in analysis_result.extracted_parties:
            print(f"- Name: {party.name}, Role: {party.role or 'N/A'}")
        print("\nDates:")
        for date_c in analysis_result.extracted_dates:
            print(f"- Type: {date_c.date_type}, Value: {date_c.date_value}, Context: {date_c.context}")
        print("\nMonetary Values:")
        for mv in analysis_result.extracted_monetary_values:
            print(f"- Amount: {mv.amount} {mv.currency}, Reason: {mv.reason or 'N/A'}")

        print("\nExtracted Clauses:")
        if analysis_result.extracted_clauses_summary:
            for clause_type, clauses in analysis_result.extracted_clauses_summary.items():
                print(f"  - {clause_type}: {len(clauses)} found")
                for clause in clauses:
                    print(f"    - Text (excerpt): {clause.get('clause_text', '')[:100]}...")
        else:
            print("  No specific clauses extracted yet.")

        print("\nAnalysis Summary (LLM Generated):")
        print(analysis_result.analysis_summary)

        print("\n--- Compliance Findings ---")
        if analysis_result.compliance_findings:
            for finding in analysis_result.compliance_findings:
                print(f"Rule ID: {finding.rule_id} ({finding.rule_description})")
                print(f"  Status: {'COMPLIANT' if finding.is_compliant else 'NON-COMPLIANT'}")
                print(f"  Severity: {finding.severity}")
                if finding.finding_details:
                    print(f"  Details: {finding.finding_details}")
                if finding.relevant_text_snippets:
                    print(f"  Relevant Snippets: {finding.relevant_text_snippets[:1]}")
                if finding.recommendation:
                    print(f"  Recommendation: {finding.recommendation}")
                print("-" * 30)
        else:
            print("  No compliance findings generated.")

        print("\n--- Knowledge Graph Summary ---")
        if analysis_result.knowledge_graph:
            print(f"  Nodes: {len(analysis_result.knowledge_graph.nodes)}")
            for node in analysis_result.knowledge_graph.nodes[:5]:
                print(f"    - ID: {node.id}, Type: {node.type}, Name: {node.name}")
            print(f"  Edges: {len(analysis_result.knowledge_graph.edges)}")
            for edge in analysis_result.knowledge_graph.edges[:5]:
                print(f"    - {edge.source_id} --({edge.type})--> {edge.target_id}")
        else:
            print("  No Knowledge Graph generated.")


        return analysis_result

    except Exception as e:
        print(f"An error occurred during document analysis for {file_path}: {e}")
        return None

async def main():
    rag_agent_instance = RAGAgent(model=llm_model, cache=redis_cache)

    dummy_docx_path = os.path.join(DOCUMENTS_DIR, "sample_agreement.docx")
    dummy_pdf_path = os.path.join(DOCUMENTS_DIR, "sample_policy.pdf")
    dummy_txt_path = os.path.join(DOCUMENTS_DIR, "sample_note.txt")

    create_dummy_docx(dummy_docx_path)
    create_dummy_pdf(dummy_pdf_path)
    create_dummy_txt(dummy_txt_path)

    analyzed_docs = []
    analyzed_docs.append(await analyze_document(dummy_docx_path, rag_agent_instance))
    analyzed_docs.append(await analyze_document(dummy_pdf_path, rag_agent_instance))
    analyzed_docs.append(await analyze_document(dummy_txt_path, rag_agent_instance))

    print("\nPhase 4 Complete: Knowledge Graph Construction and RAG Agent Integration Demonstrated.")
    print("\n--- Interactive RAG Query Session ---")
    print("You can now ask questions about the analyzed documents. Type 'exit' to quit.")

    while True:
        try:
            query = input("\nEnter your question (or 'exit'): ")
            if query.lower() == 'exit':
                break
            
            if rag_agent_instance.document_analysis_result:
                rag_response: RAGResponse = await rag_agent_instance.run(query)
                print(f"\nAI Answer (Confidence: {rag_response.confidence}):")
                print(rag_response.answer)
                if rag_response.relevant_snippets:
                    print("\nRelevant Snippets:")
                    for snippet in rag_response.relevant_snippets:
                        print(f"- {snippet}")
                if rag_response.source_nodes:
                    print("\nSource KG Nodes (IDs):")
                    for node_id in rag_response.source_nodes:
                        print(f"- {node_id}")
            else:
                print("No document has been analyzed yet. Please run the analysis first.")

        except Exception as e:
            print(f"An error occurred during RAG query: {e}")

if __name__ == "__main__":
    asyncio.run(main())
